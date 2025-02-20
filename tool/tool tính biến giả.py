#công cụ tạo biến giả
#Lưu ý, biến kiểm soát đặt ở đầu, biến phụ thuộc nằm ở cuối
import itertools
from docx import Document
#nhập dữ liệu
def nhapdulieu(kichthuoc, sochieu):
    #tạo danh sách biến
    dsb={}
    for i in range(sochieu):
        dsb[chr(65+i)]=[]
    print("nhập danh sách: ")
    i=0
    for a in dsb.keys():
        print(f"nhập biến quan sát thứ {i+1}: ")
        for j in range(kichthuoc[i]):
            bien=input()
            dsb[a].append(bien)
        i+=1
    return dsb
def taodsbienphuthuoc(dsb):
    dsbienphuthuoc=[]
    noi=itertools.product(*dsb.values())
    for i in noi:
        a=int(input(f"nhập giá trị cho {" và ".join(i)}: "))
        dsbienphuthuoc.append(a)
    return dsbienphuthuoc
def taobiengia(dsb, dsbienphuthuoc, k):
    doc = Document()  # Tạo file Word mới
    doc.add_heading("Dữ liệu Biến Giả", level=1)  # Thêm tiêu đề
    # In tiêu đề cột
    header = "\t".join(dsb.keys()) + "  <- Biến quan sát định tính"
    doc.add_paragraph(header)
    for i in dsb:
        print(i, end=" "*12)
    print("<-Biến quan sát định tính")
    #kiểm tra số lượng biến trong biến phụ thuộc
    phuthuoc=next(reversed(dsb))
    kiemtra=len(dsb[phuthuoc])
    #bắt đầu tạo biến giả
    index=[]
    noi=list(itertools.product(*dsb.values()))
    for i in range(len(dsbienphuthuoc)):
        row=[]
        for j in range(dsbienphuthuoc[i]):
            if j==0:
                for a in noi[i]:
                    for key, data_list in dsb.items():
                        if a in data_list:
                            index.append(data_list.index(a))
                index = list(map(lambda x, y: x + y, index, k ))
            print(f"{"            ".join(map(str, index))}")
            row.extend(index)
            b="\t".join(map(str, row))
            doc.add_paragraph(b)  # Ghi vào file Word
            row=[]
        index=[]
    duong_dan=r"C:\Users\LAPTOP\Desktop\tool\dulieu.docx"
    doc.save(duong_dan)
    print("Dữ liệu đã được lưu vào 'dulieu.docx' thành công!")

sochieu=int(input("nhập số chiều: "))
#bienquansat=int(input("nhập số biến quan sát: "))
kichthuoc=[]
for i in range(sochieu):
    a=int(input(f"nhập kích thước chiều thứ {i+1}: "))
    kichthuoc.append(a)
k=[]
print("nhập hệ số tăng tiến, mặc định các biến giả là 0, 1 v.v: ")#có thể nhập 0 nếu muốn giữ biến giả cơ bản hoặc nhập 1 để nhân lên thành 1, 2 v.v
for i in range(sochieu):
    r=int(input())
    k.append(r)
dulieu=nhapdulieu(kichthuoc, sochieu)
bienphuthuoc=taodsbienphuthuoc(dulieu)
taobiengia(dulieu, bienphuthuoc, k)
